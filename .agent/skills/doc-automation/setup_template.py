from docx import Document

def create_template():
    doc = Document()
    
    doc.add_heading('Lexis-AI Compliance Boutique', 0)
    
    doc.add_paragraph('Date: {{ date }}')
    
    p = doc.add_paragraph()
    p.add_run('Client: ').bold = True
    p.add_run('{{ client_name }}')
    
    p = doc.add_paragraph()
    p.add_run('Address: ').bold = True
    p.add_run('{{ company_address }}')
    
    p = doc.add_paragraph()
    p.add_run('Matter: ').bold = True
    p.add_run('{{ matter_practice_area }} - {{ matter_description }}')
    
    doc.add_heading('Engagement Letter', level=1)
    
    doc.add_paragraph(
        'Dear {{ client_name }},\n\n'
        'We are pleased to confirm our engagement to provide legal services. '
        'This letter serves as formal confirmation that your compliance onboarding is complete, '
        'including identity and background verification in accordance with the 2026 Economic Crime and Corporate Transparency Act (ECCTA).\n\n'
        'If you have any questions regarding how your data is processed, please contact our Data Protection Officer at: {{ dpo_contact }}.\n'
        'Our firm operates under the strict regulatory guidelines of the SRA (Registration: {{ sra_id }}).\n\n'
        'Please review the terms of this engagement and sign below.\n\n'
        'Sincerely,\n'
        'The Partner Team at Lexis-AI'
    )
    
    import os
    os.makedirs('templates', exist_ok=True)
    doc.save('templates/Engagement_Letter_Template.docx')
    print("Created templates/Engagement_Letter_Template.docx")

if __name__ == "__main__":
    create_template()
